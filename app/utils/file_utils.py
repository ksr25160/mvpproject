"""
Meeting AI Assistant - 파일 처리 유틸리티
"""

import streamlit as st
import tempfile
import os
import time
from datetime import datetime
import base64
import hashlib


def _create_safe_doc_id(original_filename, timestamp):
    """파일명을 Azure AI Search 호환 문서 ID로 변환합니다."""
    try:
        # 파일명에서 확장자 분리
        name_without_ext = os.path.splitext(original_filename)[0]
        extension = os.path.splitext(original_filename)[1]

        # 한글이나 특수문자가 포함된 경우 해시로 변환
        safe_name = ""
        for char in name_without_ext:
            if char.isalnum() or char in ["_", "-"]:
                safe_name += char
            else:
                # 한글이나 특수문자는 해시로 변환
                if not safe_name.endswith("_"):
                    safe_name += "_"

        # 빈 문자열이거나 너무 긴 경우 해시 사용
        if not safe_name or len(safe_name) > 50:
            # 원본 파일명의 MD5 해시 생성
            hash_name = hashlib.md5(original_filename.encode("utf-8")).hexdigest()[:12]
            safe_name = f"file_{hash_name}"

        return f"meeting_files_{safe_name}_{timestamp}"

    except Exception as e:
        # 최후의 수단: 타임스탬프만 사용
        return f"meeting_files_doc_{timestamp}"


def process_uploaded_file_from_chat(uploaded_file, service_manager, is_last_file=True):
    """파일 업로드 처리"""
    try:
        # 처리 상태 설정
        st.session_state.processing = True

        # 처리 시작 메시지 추가
        st.session_state.chat_messages.append(
            {
                "role": "assistant",
                "content": f"🔄 {uploaded_file.name} 파일을 처리하고 있습니다...",
            }
        )

        # 임시 파일로 저장
        with tempfile.NamedTemporaryFile(
            delete=False, suffix=f"_{uploaded_file.name}"
        ) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            temp_file_path = tmp_file.name

        try:
            # 파일 타입에 따른 처리
            file_extension = uploaded_file.name.lower().split(".")[-1]

            if file_extension in ["mp3", "wav", "mp4", "m4a"]:
                response = _process_audio_file(
                    uploaded_file, temp_file_path, service_manager
                )
            elif file_extension in ["pdf", "txt", "docx"]:
                response = _process_text_file(
                    uploaded_file, temp_file_path, file_extension, service_manager
                )
            else:
                response = f"❌ 지원하지 않는 파일 형식입니다: {file_extension}"

            # BlobStorage 업로드 (성공적으로 처리된 경우)
            if not response.startswith("❌"):
                try:
                    # 안전한 파일명 생성
                    timestamp = int(time.time())
                    blob_name = f"meeting_files/{uploaded_file.name}_{timestamp}"

                    # BlobStorage에 업로드
                    service_manager.upload_to_blob(temp_file_path, blob_name)

                except Exception as blob_error:
                    # BlobStorage 오류는 로그만 남기고 메인 처리는 계속
                    st.warning(f"⚠️ 파일 저장 중 오류: {str(blob_error)}")

        finally:
            # 임시 파일 삭제
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)

        # 처리 중 메시지를 결과 메시지로 업데이트
        if st.session_state.chat_messages and st.session_state.chat_messages[-1][
            "content"
        ].startswith("🔄"):
            st.session_state.chat_messages[-1]["content"] = response
        else:
            st.session_state.chat_messages.append(
                {"role": "assistant", "content": response}
            )

        # 파일 처리 완료 후 채팅 히스토리 저장 (마지막 파일에서만)
        if is_last_file:
            try:
                from utils.chat_utils import add_to_chat_history

                # 파일명을 사용자 메시지로, 처리 결과를 AI 응답으로 저장
                user_message = f"📎 파일 업로드: {uploaded_file.name}"
                add_to_chat_history(user_message, response, service_manager)

            except Exception as chat_error:
                print(f"⚠️ 채팅 히스토리 저장 오류: {str(chat_error)}")

    except Exception as e:
        error_message = f"❌ 파일 처리 중 오류가 발생했습니다: {str(e)}"

        # 처리 중 메시지를 오류 메시지로 업데이트
        if st.session_state.chat_messages and st.session_state.chat_messages[-1][
            "content"
        ].startswith("🔄"):
            st.session_state.chat_messages[-1]["content"] = error_message
        else:
            st.session_state.chat_messages.append(
                {"role": "assistant", "content": error_message}
            )
    finally:
        # 처리 상태 해제 (마지막 파일에서만)
        if is_last_file:
            st.session_state.processing = False
        # 파일 처리 완료 후 UI 업데이트 (마지막 파일에서만)
        if is_last_file:
            st.rerun()


def _process_audio_file(uploaded_file, temp_file_path, service_manager):
    """음성 파일 처리"""
    try:
        transcribed_text = service_manager.transcribe_audio(temp_file_path)

        if transcribed_text:
            # AI 분석 먼저 수행
            analysis_result = service_manager.summarize_and_extract(transcribed_text)

            # STT 텍스트를 별도 파일로 Blob Storage에 저장
            try:
                # STT 텍스트 파일 이름 생성
                stt_blob_name = (
                    f"meeting_files/{uploaded_file.name}_stt_{int(time.time())}.txt"
                )

                # 임시 STT 텍스트 파일 생성
                import tempfile

                with tempfile.NamedTemporaryFile(
                    mode="w", encoding="utf-8", suffix=".txt", delete=False
                ) as stt_file:
                    stt_file.write(transcribed_text)
                    stt_temp_path = stt_file.name

                # STT 텍스트를 Blob Storage에 업로드
                service_manager.upload_to_blob(stt_temp_path, stt_blob_name)

                # STT 텍스트를 AI Search에도 인덱싱 (AI 분석 결과 포함)
                try:
                    # 안전한 문서 ID 생성
                    timestamp = int(time.time())
                    safe_doc_id = _create_safe_doc_id(
                        f"{uploaded_file.name}_stt", timestamp
                    )

                    # AI 분석 결과를 포함한 메타데이터
                    metadata = {
                        "filename": f"{uploaded_file.name} (STT)",
                        "upload_time": datetime.now().isoformat(),
                        "file_type": "stt_audio",
                        "original_file": uploaded_file.name,
                        "meeting_title": analysis_result.get("meetingTitle", ""),
                        "participants": ", ".join(
                            analysis_result.get("participants", [])
                        ),
                        "summary": analysis_result.get("summary", ""),
                    }
                    service_manager.index_document(
                        doc_id=safe_doc_id,
                        content=transcribed_text,
                        metadata=metadata,
                        blob_path=stt_blob_name,
                    )
                    print(
                        f"✅ STT 텍스트 AI Search 인덱싱 완료 (참석자 정보 포함): {safe_doc_id}"
                    )
                except Exception as index_error:
                    print(f"⚠️ STT 텍스트 AI Search 인덱싱 실패: {str(index_error)}")

                # 임시 파일 삭제
                import os

                os.unlink(stt_temp_path)

                print(f"✅ STT 텍스트 Blob Storage 저장 완료: {stt_blob_name}")

            except Exception as stt_blob_error:
                print(f"⚠️ STT 텍스트 Blob Storage 저장 실패: {str(stt_blob_error)}")

            # 액션 아이템에 담당자 추천 추가
            if analysis_result.get("actionItems"):
                for action_item in analysis_result["actionItems"]:
                    task_description = action_item.get("description", "")
                    if task_description:
                        try:
                            # RAG 기반 담당자 검색 및 추천
                            from services.search_service import search_staff_for_task

                            staff_candidates = search_staff_for_task(
                                task_description, top_k=3
                            )

                            if staff_candidates:
                                # 가장 적합한 담당자 선택
                                best_candidate = staff_candidates[0]
                                action_item["assignee"] = best_candidate["name"]
                                action_item["recommended_department"] = best_candidate[
                                    "department"
                                ]
                            else:
                                # RAG 검색 실패 시 미할당 처리
                                action_item["assignee"] = "미할당"
                        except Exception as e:
                            print(f"⚠️ 담당자 추천 오류: {str(e)}")
                            action_item["assignee"] = "담당자 미정"

            # 결과를 Cosmos DB에 저장
            meeting_id = service_manager.save_meeting(
                meeting_title=analysis_result.get("meetingTitle", uploaded_file.name),
                raw_text=transcribed_text,
                summary_json=analysis_result,
            )

            # 응답 메시지 생성
            response = f"""✅ **음성 파일 분석 완료**

**파일명:** {uploaded_file.name}
**회의 제목:** {analysis_result.get('meetingTitle', 'N/A')}
**참석자:** {', '.join(analysis_result.get('participants', ['N/A']))}

**요약:**
{analysis_result.get('summary', 'N/A')}

**액션 아이템:** {len(analysis_result.get('actionItems', []))}개"""

            # 액션 아이템 목록을 담당자와 함께 표시
            for i, item in enumerate(analysis_result.get("actionItems", [])[:5], 1):
                assignee = item.get("assignee", "N/A")
                response += (
                    f"\n{i}. {item.get('description', 'N/A')} (담당: {assignee})"
                )

            if len(analysis_result.get("actionItems", [])) > 5:
                response += (
                    f"\n... 외 {len(analysis_result.get('actionItems', [])) - 5}개"
                )

        else:
            response = f"❌ 음성 파일 {uploaded_file.name}의 전사에 실패했습니다."

        return response

    except Exception as e:
        return f"❌ 음성 파일 처리 중 오류가 발생했습니다: {str(e)}"


def _process_text_file(uploaded_file, temp_file_path, file_extension, service_manager):
    """텍스트 파일 처리"""
    try:
        if file_extension == "txt":
            with open(temp_file_path, "r", encoding="utf-8") as f:
                file_content = f.read()
        elif file_extension == "docx":
            # DOCX 파일 파싱
            try:
                from docx import Document

                doc = Document(temp_file_path)
                file_content = ""
                for paragraph in doc.paragraphs:
                    file_content += paragraph.text + "\n"

                # 표(table) 내용도 추출
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            file_content += cell.text + " "
                        file_content += "\n"

            except ImportError:
                return "❌ DOCX 파일 처리를 위해 python-docx 패키지가 필요합니다. 설치 후 다시 시도해주세요."
            except Exception as e:
                return f"❌ DOCX 파일 파싱 중 오류가 발생했습니다: {str(e)}"

        elif file_extension == "pdf":
            # PDF 파일 파싱
            try:
                import PyPDF2

                with open(temp_file_path, "rb") as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    file_content = ""
                    for page in pdf_reader.pages:
                        file_content += page.extract_text() + "\n"
            except ImportError:
                return "❌ PDF 파일 처리를 위해 PyPDF2 패키지가 필요합니다. 설치 후 다시 시도해주세요."
            except Exception as e:
                return f"❌ PDF 파일 파싱 중 오류가 발생했습니다: {str(e)}"
        else:
            return f"❌ 지원하지 않는 파일 형식입니다: {file_extension}"

        if file_content:
            # AI 분석
            analysis_result = service_manager.summarize_and_extract(file_content)

            # AI Search에 인덱싱 (AI 분석 결과 포함)
            try:
                # 안전한 문서 ID 생성
                timestamp = int(time.time())
                safe_doc_id = _create_safe_doc_id(uploaded_file.name, timestamp)
                blob_name = f"meeting_files/{uploaded_file.name}_{timestamp}"

                # AI 분석 결과를 포함한 메타데이터
                metadata = {
                    "filename": uploaded_file.name,
                    "upload_time": datetime.now().isoformat(),
                    "file_type": file_extension,
                    "meeting_title": analysis_result.get("meetingTitle", ""),
                    "participants": ", ".join(analysis_result.get("participants", [])),
                    "summary": analysis_result.get("summary", ""),
                }
                service_manager.index_document(
                    doc_id=safe_doc_id,
                    content=file_content,
                    metadata=metadata,
                    blob_path=blob_name,
                )
                print(
                    f"✅ 문서 AI Search 인덱싱 완료 (참석자 정보 포함): {safe_doc_id}"
                )
            except Exception as index_error:
                print(f"⚠️ 문서 AI Search 인덱싱 실패: {str(index_error)}")

            # 액션 아이템에 담당자 추천 추가
            if analysis_result.get("actionItems"):
                for action_item in analysis_result["actionItems"]:
                    task_description = action_item.get("description", "")
                    if task_description:
                        try:
                            # RAG 기반 담당자 검색 및 추천
                            from services.search_service import search_staff_for_task

                            staff_candidates = search_staff_for_task(
                                task_description, top_k=3
                            )

                            if staff_candidates:
                                # 가장 적합한 담당자 선택
                                best_candidate = staff_candidates[0]
                                action_item["assignee"] = best_candidate["name"]
                                action_item["recommended_department"] = best_candidate[
                                    "department"
                                ]
                            else:
                                # RAG 검색 실패 시 미할당 처리
                                action_item["assignee"] = "미할당"
                        except Exception as e:
                            print(f"⚠️ 담당자 추천 오류: {str(e)}")
                            action_item["assignee"] = "담당자 미정"

            # 결과를 Cosmos DB에 저장
            meeting_id = service_manager.save_meeting(
                meeting_title=analysis_result.get("meetingTitle", uploaded_file.name),
                raw_text=file_content,
                summary_json=analysis_result,
            )

            response = f"""✅ **문서 분석 완료**

**파일명:** {uploaded_file.name}
**회의 제목:** {analysis_result.get('meetingTitle', 'N/A')}
**참석자:** {', '.join(analysis_result.get('participants', ['N/A']))}

**요약:**
{analysis_result.get('summary', 'N/A')}

**액션 아이템:** {len(analysis_result.get('actionItems', []))}개"""

            # 액션 아이템 목록을 담당자와 함께 표시
            for i, item in enumerate(analysis_result.get("actionItems", [])[:5], 1):
                assignee = item.get("assignee", "N/A")
                response += (
                    f"\n{i}. {item.get('description', 'N/A')} (담당: {assignee})"
                )

            if len(analysis_result.get("actionItems", [])) > 5:
                response += (
                    f"\n... 외 {len(analysis_result.get('actionItems', [])) - 5}개"
                )
        else:
            response = f"❌ 파일 {uploaded_file.name}을 읽을 수 없습니다."

        return response

    except Exception as e:
        return f"❌ 파일 처리 중 오류가 발생했습니다: {str(e)}"
